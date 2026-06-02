#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成 Codex APP 完整说明 + 手动操作清单 Word 文档
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── 颜色常量 ──────────────────────────────────────────
C_BLUE   = RGBColor(0x1F, 0x49, 0x7D)   # 深蓝，标题
C_GREEN  = RGBColor(0x1D, 0x6A, 0x39)   # 深绿，已完成
C_AMBER  = RGBColor(0x7B, 0x3F, 0x00)   # 深橙，手动操作
C_LGRAY  = RGBColor(0xF2, 0xF2, 0xF2)   # 浅灰，表格底色
C_DKGRAY = RGBColor(0x40, 0x40, 0x40)   # 深灰，正文

def set_cell_bg(cell, hex_color: str):
    """设置单元格背景色"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_row_bg(row, hex_color: str):
    for cell in row.cells:
        set_cell_bg(cell, hex_color)

def set_para_space(para, before=0, after=80):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(before))
    spacing.set(qn('w:after'), str(after))
    pPr.append(spacing)

def add_heading_styled(doc, text, level=1, color=None):
    """自定义颜色标题"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(16)
    elif level == 2:
        run.font.size = Pt(13)
    else:
        run.font.size = Pt(11.5)
    run.font.color.rgb = color or C_BLUE
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    set_para_space(p, before=120, after=60)
    return p

def add_status_badge(para, text, color):
    """在段落中添加状态徽章"""
    run = para.add_run(f'  [{text}]')
    run.bold = True
    run.font.color.rgb = color
    run.font.size = Pt(10)

def make_table_header(table, headers, bg='1F497D'):
    """设置表头行"""
    row = table.rows[0]
    set_row_bg(row, bg)
    for i, h in enumerate(headers):
        cell = row.cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.font.size = Pt(10.5)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_table_row(table, row_idx, data, bg=None, bold_first=False):
    row = table.rows[row_idx]
    if bg:
        set_row_bg(row, bg)
    for i, val in enumerate(data):
        cell = row.cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(str(val))
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.font.size = Pt(10)
        if bold_first and i == 0:
            run.bold = True


def build_doc():
    doc = Document()

    # ── 页面设置：A4 ───────────────────────────────────
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width  = Cm(21.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.0)

    # 默认正文样式
    normal = doc.styles['Normal']
    normal.font.name = '微软雅黑'
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # ══════════════════════════════════════════════════
    # 封面标题
    # ══════════════════════════════════════════════════
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_space(title_p, before=200, after=40)
    r = title_p.add_run('Codex APP 完整安装说明\n& 手动操作清单')
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = C_BLUE
    r.font.name = '微软雅黑'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_space(sub_p, before=0, after=200)
    sr = sub_p.add_run('WorkBuddy AI 自动生成 · 2026-05-03\n合肥工业大学（宣城校区）机械工程专业')
    sr.font.size = Pt(10)
    sr.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    sr.font.name = '微软雅黑'
    sr._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    doc.add_paragraph()

    # ══════════════════════════════════════════════════
    # 第一部分：项目进度总览
    # ══════════════════════════════════════════════════
    add_heading_styled(doc, '一、项目进度总览', 1)

    overview_p = doc.add_paragraph()
    overview_p.add_run('AI 已自动完成所有可自动化步骤。你只需完成 ').font.size = Pt(11)
    rb = overview_p.add_run('3 项手动操作')
    rb.bold = True
    rb.font.size = Pt(11)
    rb.font.color.rgb = C_AMBER
    overview_p.add_run('（约 10–15 分钟）即可开始使用 Codex APP。').font.size = Pt(11)
    set_para_space(overview_p, after=120)

    # 总览表格
    tbl = doc.add_table(rows=8, cols=4)
    tbl.style = 'Table Grid'
    make_table_header(tbl, ['类别', '项目', '状态', '操作人'], '1F497D')

    rows_data = [
        ('安装', 'Node.js v22.15.0',          '✅ 已完成', 'AI'),
        ('安装', 'Git v2.47.1',               '✅ 已完成', 'AI'),
        ('安装', 'VSCode v1.118.1',            '✅ 已完成', 'AI'),
        ('安装', 'Codex APP v24.11.1',         '✅ 已完成', 'AI'),
        ('配置', 'AGENTS.md 全局配置',         '✅ 已完成', 'AI'),
        ('配置', 'config.toml（IDE=VSCode）',  '✅ 已完成', 'AI'),
        ('使用', '登录 / 沙箱 / GitHub',       '⚠️ 待用户操作', '你'),
    ]

    bg_colors = ['FFFFFF', 'F2F2F2'] * 10
    for i, (cat, item, status, operator) in enumerate(rows_data, 1):
        bg = bg_colors[i % 2]
        add_table_row(tbl, i, [cat, item, status, operator], bg=bg)
        # 状态列颜色
        cell_s = tbl.rows[i].cells[2]
        run_s = cell_s.paragraphs[0].runs[0]
        if '已完成' in status:
            run_s.font.color.rgb = C_GREEN
        else:
            run_s.font.color.rgb = C_AMBER
        run_s.bold = True
        # 操作人颜色
        cell_o = tbl.rows[i].cells[3]
        run_o = cell_o.paragraphs[0].runs[0]
        if operator == '你':
            run_o.font.color.rgb = C_AMBER
            run_o.bold = True

    doc.add_paragraph()

    # ══════════════════════════════════════════════════
    # 第二部分：已安装组件详情
    # ══════════════════════════════════════════════════
    add_heading_styled(doc, '二、已安装组件详情（AI 已完成）', 1, C_GREEN)

    # 工具路径表
    tbl2 = doc.add_table(rows=5, cols=3)
    tbl2.style = 'Table Grid'
    make_table_header(tbl2, ['工具', '版本', '安装路径'], '1D6A39')

    tools = [
        ('Node.js',    'v22.15.0',  r'C:\Program Files\nodejs\node.exe'),
        ('Git',        'v2.47.1',   r'C:\Program Files\Git\cmd\git.exe'),
        ('VSCode',     'v1.118.1',  r'D:\Program Files\Microsoft VS Code\Code.exe'),
        ('Codex APP',  'v24.11.1',  r'C:\Users\Matebook\AppData\Local\Codex\Codex.exe'),
    ]
    for i, row in enumerate(tools, 1):
        add_table_row(tbl2, i, row, bg='FFFFFF' if i % 2 == 1 else 'F0FAF4', bold_first=True)

    doc.add_paragraph()

    # 配置文件表
    p = doc.add_paragraph()
    p.add_run('📁 配置文件位置').bold = True

    tbl3 = doc.add_table(rows=3, cols=3)
    tbl3.style = 'Table Grid'
    make_table_header(tbl3, ['文件', '路径', '说明'], '1D6A39')

    cfg_data = [
        ('AGENTS.md',    r'C:\Users\Matebook\.codex\AGENTS.md',    '全局 AI 配置（用户背景 + 安全规则）'),
        ('config.toml',  r'C:\Users\Matebook\.codex\config.toml',  'IDE=VSCode，沙箱=windows-native'),
    ]
    for i, row in enumerate(cfg_data, 1):
        add_table_row(tbl3, i, row, bg='FFFFFF' if i % 2 == 1 else 'F0FAF4')

    doc.add_paragraph()

    # ══════════════════════════════════════════════════
    # 第三部分：手动操作清单（重点）
    # ══════════════════════════════════════════════════
    add_heading_styled(doc, '三、⚠️ 手动操作清单（你需要完成）', 1, C_AMBER)

    intro = doc.add_paragraph()
    intro.add_run('下列步骤 AI 无法代劳（需要你的账号 / 点击确认），请按顺序完成。').italic = True
    intro.add_run(' 预计总耗时：').font.size = Pt(10.5)
    rb2 = intro.add_run('10–15 分钟')
    rb2.bold = True
    rb2.font.color.rgb = C_AMBER
    set_para_space(intro, after=100)

    # ── 手动步骤详情表 ────────────────────────────────
    manual_steps = [
        (
            '步骤 1', '必须', '启动 Codex APP',
            '双击桌面图标  或  运行：d:\\杂物间\\start_codex.bat\n'
            '（start_codex.bat 会同时打开 VSCode + Codex）',
            '约 1 分钟',
        ),
        (
            '步骤 2', '必须', '登录 ChatGPT 账号',
            'Codex 界面 → 点击 "Login" / "Sign in"\n'
            '→ 使用你的 ChatGPT 账号（需要 ChatGPT Plus 订阅）',
            '约 2 分钟',
        ),
        (
            '步骤 3', '必须', '初始化沙箱',
            '登录后 → 点击 "设置沙箱" / "Setup Sandbox"\n'
            '→ 等待初始化完成（首次需下载镜像）',
            '约 5–10 分钟',
        ),
        (
            '步骤 4', '可选', '配置 GitHub（代码推送）',
            '① 生成 SSH 密钥：\n'
            '   ssh-keygen -t ed25519 -C "your_email"\n'
            '② 将公钥添加到 GitHub Settings → SSH keys\n'
            '③ Codex → Settings → Integrations → GitHub → Connect',
            '约 5 分钟',
        ),
        (
            '步骤 5', '可选', '安装 Netlify / MCP 插件',
            'Codex → Settings → Plugins\n'
            '→ 搜索并安装 Netlify（部署网页）或其他 MCP 插件',
            '按需',
        ),
    ]

    tbl_m = doc.add_table(rows=len(manual_steps)+1, cols=5)
    tbl_m.style = 'Table Grid'
    make_table_header(tbl_m, ['步骤', '类型', '任务', '具体操作', '预计时间'], '7B3F00')

    for i, (step, stype, task, detail, time) in enumerate(manual_steps, 1):
        row = tbl_m.rows[i]
        bg = 'FFF8F0' if i % 2 == 1 else 'FFFFFF'
        set_row_bg(row, bg)

        # 步骤号
        c0 = row.cells[0]
        c0.text = ''
        r0 = c0.paragraphs[0].add_run(step)
        r0.bold = True
        r0.font.color.rgb = C_BLUE
        r0.font.name = '微软雅黑'
        r0._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        r0.font.size = Pt(10)
        c0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 类型
        c1 = row.cells[1]
        c1.text = ''
        r1 = c1.paragraphs[0].add_run(stype)
        r1.bold = True
        r1.font.color.rgb = C_AMBER if stype == '必须' else RGBColor(0x40, 0x70, 0x40)
        r1.font.name = '微软雅黑'
        r1._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        r1.font.size = Pt(10)
        c1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 任务名
        c2 = row.cells[2]
        c2.text = ''
        r2 = c2.paragraphs[0].add_run(task)
        r2.bold = True
        r2.font.name = '微软雅黑'
        r2._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        r2.font.size = Pt(10)

        # 操作详情（换行处理）
        c3 = row.cells[3]
        for j, line in enumerate(detail.split('\n')):
            if j == 0:
                p3 = c3.paragraphs[0]
            else:
                p3 = c3.add_paragraph()
            r3 = p3.add_run(line)
            r3.font.name = '微软雅黑'
            r3._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            r3.font.size = Pt(9.5)
            if line.startswith('   ') or line.startswith('→'):
                r3.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

        # 时间
        c4 = row.cells[4]
        c4.text = ''
        r4 = c4.paragraphs[0].add_run(time)
        r4.font.name = '微软雅黑'
        r4._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        r4.font.size = Pt(10)
        c4.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # ── 手动操作 Checklist ────────────────────────────
    add_heading_styled(doc, '✅ 手动操作 Checklist（打勾确认）', 2, C_AMBER)

    checklist = [
        ('□', '必须', '启动 Codex APP（运行 start_codex.bat 或桌面图标）'),
        ('□', '必须', '登录 ChatGPT Plus 账号'),
        ('□', '必须', '初始化沙箱（Setup Sandbox，等待完成）'),
        ('□', '可选', '配置 GitHub SSH + Codex 授权'),
        ('□', '可选', '安装 Netlify / MCP 等插件'),
    ]

    for box, stype, desc in checklist:
        p = doc.add_paragraph()
        set_para_space(p, before=40, after=40)
        r_box = p.add_run(f'  {box}  ')
        r_box.font.size = Pt(14)
        r_box.font.color.rgb = C_AMBER if stype == '必须' else RGBColor(0x40, 0x70, 0x40)

        r_type = p.add_run(f'[{stype}]  ')
        r_type.bold = True
        r_type.font.size = Pt(10)
        r_type.font.color.rgb = C_AMBER if stype == '必须' else RGBColor(0x40, 0x70, 0x40)
        r_type.font.name = '微软雅黑'
        r_type._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

        r_desc = p.add_run(desc)
        r_desc.font.size = Pt(10.5)
        r_desc.font.name = '微软雅黑'
        r_desc._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    doc.add_paragraph()

    # ══════════════════════════════════════════════════
    # 第四部分：常用命令速查
    # ══════════════════════════════════════════════════
    add_heading_styled(doc, '四、常用命令速查', 1)

    cmds = [
        ('启动 Codex（批处理）',  r'd:\杂物间\start_codex.bat',              '双击运行'),
        ('验证环境',              r'd:\杂物间\verify_codex_env.bat',          '双击运行'),
        ('验证 Node.js',         r'node --version',                          '在终端执行'),
        ('验证 Git',             r'git --version',                           '在终端执行'),
        ('验证 VSCode',          r'code --version',                          '在终端执行'),
        ('生成 SSH 密钥',        r'ssh-keygen -t ed25519 -C "你的邮箱"',     '在 Git Bash 执行'),
        ('查看 SSH 公钥',        r'cat ~/.ssh/id_ed25519.pub',               '在 Git Bash 执行'),
    ]

    tbl4 = doc.add_table(rows=len(cmds)+1, cols=3)
    tbl4.style = 'Table Grid'
    make_table_header(tbl4, ['操作', '命令 / 路径', '说明'], '1F497D')

    for i, (op, cmd, note) in enumerate(cmds, 1):
        row = tbl4.rows[i]
        bg = 'FFFFFF' if i % 2 == 1 else 'F2F6FF'
        set_row_bg(row, bg)

        row.cells[0].text = ''
        r_op = row.cells[0].paragraphs[0].add_run(op)
        r_op.font.name = '微软雅黑'
        r_op._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        r_op.font.size = Pt(10)

        row.cells[1].text = ''
        r_cmd = row.cells[1].paragraphs[0].add_run(cmd)
        r_cmd.font.name = 'Courier New'
        r_cmd.font.size = Pt(9.5)
        r_cmd.font.color.rgb = RGBColor(0x1A, 0x1A, 0x6E)

        row.cells[2].text = ''
        r_note = row.cells[2].paragraphs[0].add_run(note)
        r_note.font.name = '微软雅黑'
        r_note._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        r_note.font.size = Pt(10)
        r_note.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    doc.add_paragraph()

    # ══════════════════════════════════════════════════
    # 第五部分：故障排除
    # ══════════════════════════════════════════════════
    add_heading_styled(doc, '五、常见问题排除', 1)

    issues = [
        ('命令未找到（node/git/code）',
         '原因：新安装后需重启终端，环境变量才生效。\n'
         '解决：关闭当前终端，重新打开；或运行 verify_codex_env.bat 查看状态。'),
        ('PowerShell 提示"无法加载 npm.ps1"',
         '原因：执行策略限制。\n'
         '解决：改用 npm.cmd 替代 npm；或以管理员身份执行：\n'
         '  Set-ExecutionPolicy -ExecutionPolicy RemoteSigned -Scope CurrentUser'),
        ('Codex 无法启动沙箱',
         '原因：Windows Sandbox 功能未启用。\n'
         '解决：控制面板 → 程序 → 启用或关闭 Windows 功能 → 勾选"Windows 沙盒" → 重启。'),
        ('登录 ChatGPT 失败',
         '原因：网络访问受限，或账号未开通 Plus。\n'
         '解决：检查网络连接；确认 ChatGPT 账号为 Plus 或以上订阅。'),
    ]

    for q, a in issues:
        p_q = doc.add_paragraph()
        set_para_space(p_q, before=80, after=20)
        rq = p_q.add_run(f'Q: {q}')
        rq.bold = True
        rq.font.color.rgb = C_BLUE
        rq.font.name = '微软雅黑'
        rq._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        rq.font.size = Pt(10.5)

        for line in a.split('\n'):
            p_a = doc.add_paragraph()
            set_para_space(p_a, before=0, after=20)
            ra = p_a.add_run(f'    {line}')
            ra.font.name = '微软雅黑'
            ra._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            ra.font.size = Pt(10)
            ra.font.color.rgb = C_DKGRAY

    doc.add_paragraph()

    # ══════════════════════════════════════════════════
    # 页脚：文件位置说明
    # ══════════════════════════════════════════════════
    add_heading_styled(doc, '附：所有生成文件位置', 2, RGBColor(0x40, 0x40, 0x40))

    files = [
        ('本文档',               r'd:\杂物间\Codex_APP_说明与手动清单.docx'),
        ('16 章节全攻略说明书',  r'd:\杂物间\Codex_APP_保姆级全攻略_说明书.md'),
        ('安装配置指南',         r'd:\杂物间\Codex_APP_完整安装配置指南.md'),
        ('环境验证脚本',         r'd:\杂物间\verify_codex_env.bat'),
        ('快速启动脚本',         r'd:\杂物间\start_codex.bat'),
        ('Codex 全局配置',       r'C:\Users\Matebook\.codex\AGENTS.md'),
    ]

    tbl5 = doc.add_table(rows=len(files)+1, cols=2)
    tbl5.style = 'Table Grid'
    make_table_header(tbl5, ['文件说明', '路径'], '404040')

    for i, (desc, path) in enumerate(files, 1):
        row = tbl5.rows[i]
        set_row_bg(row, 'FFFFFF' if i % 2 == 1 else 'F5F5F5')
        row.cells[0].text = ''
        rr = row.cells[0].paragraphs[0].add_run(desc)
        rr.font.name = '微软雅黑'
        rr._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        rr.font.size = Pt(10)

        row.cells[1].text = ''
        rp = row.cells[1].paragraphs[0].add_run(path)
        rp.font.name = 'Courier New'
        rp.font.size = Pt(9)
        rp.font.color.rgb = RGBColor(0x1A, 0x1A, 0x6E)

    # 结尾语
    p_end = doc.add_paragraph()
    set_para_space(p_end, before=200, after=0)
    r_end = p_end.add_run('按照手动操作清单完成 3 项必须步骤后，即可正式使用 Codex APP。祝你顺利！🚀')
    r_end.bold = True
    r_end.font.size = Pt(11)
    r_end.font.color.rgb = C_BLUE
    r_end.font.name = '微软雅黑'
    r_end._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 保存
    out = r'd:\杂物间\Codex_APP_说明与手动清单.docx'
    doc.save(out)
    print(f'✓ 已生成：{out}')


if __name__ == '__main__':
    build_doc()
