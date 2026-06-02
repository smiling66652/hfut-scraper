#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Markdown 到 Word 转换器
将 Codex APP 项目交付报告转换为 Word 文档
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

def create_document():
    """创建 Word 文档"""
    doc = Document()
    
    # 设置默认字体（中文）
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(10.5)
    
    # 添加标题
    title = doc.add_heading('Codex APP 项目复刻 - 最终交付报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加副标题信息
    p = doc.add_paragraph()
    p.add_run('项目时间：2026-05-02 至 2026-05-03\n').bold = True
    p.add_run('执行者：WorkBuddy AI\n')
    p.add_run('用户：合肥工业大学（宣城校区）机械工程专业')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # 项目概述
    doc.add_heading('📋 项目概述', 1)
    doc.add_paragraph('根据 B站教程（https://www.bilibili.com/opus/1196732919815602198），成功复刻 Codex APP 开发环境，完成所有可自动化步骤。')
    
    # 已完成部分
    doc.add_heading('✅ 已完成（AI 自动执行）', 1)
    
    # 1. 环境安装
    doc.add_heading('1. 环境安装', 2)
    table = doc.add_table(rows=6, cols=4)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = '组件'
    hdr[1].text = '版本'
    hdr[2].text = '安装路径'
    hdr[3].text = '状态'
    
    data = [
        ('Node.js', 'v22.15.0', 'C:\\Program Files\\nodejs\\', '✅ 已完成'),
        ('npm', '10.9.2', 'C:\\Program Files\\nodejs\\', '✅ 已完成'),
        ('Git', 'v2.47.1', 'C:\\Program Files\\Git\\', '✅ 已完成'),
        ('VSCode', 'v1.118.1', 'D:\\Program Files\\Microsoft VS Code\\', '✅ 已完成'),
        ('Codex APP', 'v24.11.1', 'C:\\Users\\Matebook\\AppData\\Local\\Codex\\', '✅ 已完成'),
    ]
    
    for i, (comp, ver, path, status) in enumerate(data, 1):
        cells = table.rows[i].cells
        cells[0].text = comp
        cells[1].text = ver
        cells[2].text = path
        cells[3].text = status
    
    doc.add_paragraph()
    
    # 2. 配置文件
    doc.add_heading('2. 配置文件', 2)
    table2 = doc.add_table(rows=3, cols=3)
    table2.style = 'Light Grid Accent 1'
    hdr2 = table2.rows[0].cells
    hdr2[0].text = '文件'
    hdr2[1].text = '路径'
    hdr2[2].text = '说明'
    
    configs = [
        ('AGENTS.md', 'C:\\Users\\Matebook\\.codex\\AGENTS.md', '全局 AI 配置（含用户背景、安全规则）'),
        ('config.toml', 'C:\\Users\\Matebook\\.codex\\config.toml', 'Codex 配置文件（IDE 已设为 VSCode）'),
    ]
    
    for i, (file, path, desc) in enumerate(configs, 1):
        cells = table2.rows[i].cells
        cells[0].text = file
        cells[1].text = path
        cells[2].text = desc
    
    doc.add_paragraph()
    
    # 3. 辅助脚本
    doc.add_heading('3. 辅助脚本', 2)
    table3 = doc.add_table(rows=4, cols=3)
    table3.style = 'Light Grid Accent 1'
    hdr3 = table3.rows[0].cells
    hdr3[0].text = '脚本'
    hdr3[1].text = '路径'
    hdr3[2].text = '用途'
    
    scripts = [
        ('verify_codex_env.sh', 'd:\\杂物间\\verify_codex_env.sh', 'Bash 环境验证脚本'),
        ('verify_codex_env.bat', 'd:\\杂物间\\verify_codex_env.bat', 'Windows 环境验证脚本'),
        ('start_codex.bat', 'd:\\杂物间\\start_codex.bat', 'Codex 快速启动脚本'),
    ]
    
    for i, (script, path, desc) in enumerate(scripts, 1):
        cells = table3.rows[i].cells
        cells[0].text = script
        cells[1].text = path
        cells[2].text = desc
    
    doc.add_paragraph()
    
    # 4. 文档生成
    doc.add_heading('4. 文档生成', 2)
    table4 = doc.add_table(rows=4, cols=3)  # 1 header + 3 data rows
    table4.style = 'Light Grid Accent 1'
    hdr4 = table4.rows[0].cells
    hdr4[0].text = '文档'
    hdr4[1].text = '路径'
    hdr4[2].text = '说明'
    
    docs = [
        ('Codex_APP_保姆级全攻略_说明书.md', 'd:\\杂物间\\', '16章节完整教程'),
        ('Codex_APP_完整安装配置指南.md', 'd:\\杂物间\\', '安装配置快速指南'),
        ('Codex_APP_项目交付报告.md', 'd:\\杂物间\\', '本项目交付报告'),
    ]
    
    for i, (doc_name, path, desc) in enumerate(docs, 1):
        cells = table4.rows[i].cells
        cells[0].text = doc_name
        cells[1].text = path
        cells[2].text = desc
    
    doc.add_paragraph()
    
    # 需用户手动完成部分
    doc.add_heading('⚠️ 需用户手动完成', 1)
    
    steps = [
        ('第一步：启动并登录 Codex', [
            '启动 Codex APP（双击桌面图标，或开始菜单搜索，或运行 start_codex.bat）',
            '登录 ChatGPT 账号（需要 Plus 订阅）',
        ]),
        ('第二步：初始化沙箱', [
            '登录后，点击界面中的"设置沙箱"或"Setup Sandbox"',
            '等待初始化完成（需要下载沙箱镜像，约 5-10 分钟）',
        ]),
        ('第三步：配置 IDE（已自动配置）', [
            '已在 config.toml 中设置 ide = "vscode"',
            '如需手动确认：Codex 界面 → Settings → Default IDE → 选择 "Visual Studio Code"',
        ]),
        ('第四步（可选）：配置 GitHub', [
            '生成 SSH 密钥：ssh-keygen -t ed25519 -C "your_email@example.com"',
            '添加 SSH 密钥到 GitHub（Settings → SSH and GPG keys）',
            '授权 Codex 访问 GitHub（Settings → Integrations → GitHub → Connect）',
        ]),
        ('第五步（可选）：安装插件', [
            'Netlify 插件：用于部署网页应用',
            'MCP 服务器：用于扩展功能（如数据库、API 等）',
        ]),
    ]
    
    for step_title, step_items in steps:
        doc.add_heading(step_title, 2)
        for item in step_items:
            doc.add_paragraph(item, style='List Bullet')
    
    # 验证安装
    doc.add_heading('🔧 验证安装', 1)
    
    doc.add_heading('方法1：运行验证脚本', 2)
    p = doc.add_paragraph()
    p.add_run('Windows 用户：\n').bold = True
    p.add_run('  d:\\杂物间\\verify_codex_env.bat\n\n')
    p.add_run('Git Bash 用户：\n').bold = True
    p.add_run('  bash /d/杂物间/verify_codex_env.sh')
    
    doc.add_heading('方法2：手动验证', 2)
    p = doc.add_paragraph('打开新的终端（PowerShell 或 Git Bash），执行：\n')
    p.add_run('node --version\n').bold = True
    p.add_run('npm --version\n')
    p.add_run('git --version\n')
    p.add_run('code --version\n').bold = True
    
    p = doc.add_paragraph()
    p.add_run('如果提示"命令未找到"，请重启终端或执行：\n').italic = True
    p.add_run('export PATH="$PATH:/c/Program Files/nodejs:/c/Program Files/Git/cmd:/d/Program Files/Microsoft VS Code/bin"').italic = True
    
    # 使用流程
    doc.add_heading('📝 使用流程', 1)
    
    usage = [
        ('基本使用', [
            '启动 Codex（运行 start_codex.bat 或直接启动 Codex APP）',
            '在对话框中描述你想完成的任务（例如："帮我创建一个 Python 爬虫，抓取天气预报"）',
            'Codex 会生成代码并在 VSCode 中打开',
            '审查代码，确认无误后点击 "Run" 执行',
        ]),
        ('高级功能', [
            '沙箱测试：在隔离环境中安全运行代码',
            '多文件编辑：同时编辑多个文件',
            '版本控制：自动 git commit',
            '插件扩展：连接数据库、API 等外部服务',
        ]),
    ]
    
    for usage_title, usage_items in usage:
        doc.add_heading(usage_title, 2)
        for i, item in enumerate(usage_items, 1):
            p = doc.add_paragraph()
            p.add_run(f'{i}. ').bold = True
            p.add_run(item)
    
    # 总结
    doc.add_heading('🎯 总结', 1)
    
    p = doc.add_paragraph()
    p.add_run('✅ 已完成：').bold = True
    p.add_run('5 个工具安装，2 个配置文件，3 个辅助脚本，2 份文档\n')
    
    p = doc.add_paragraph()
    p.add_run('⚠️ 需手动：').bold = True
    p.add_run('登录 ChatGPT，初始化沙箱，配置 GitHub（可选）\n')
    
    p = doc.add_paragraph()
    p.add_run('⏱️ 预计耗时：').bold = True
    p.add_run('手动操作约 10-15 分钟')
    
    # 添加结尾
    p = doc.add_paragraph()
    p.add_run('\n祝你使用愉快！🚀').bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    p.add_run('本文档由 WorkBuddy AI 自动生成 - 2026-05-03')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # 保存文档
    output_path = 'd:/杂物间/Codex_APP_项目交付报告.docx'
    doc.save(output_path)
    print(f'✓ 文档已生成：{output_path}')
    return output_path

if __name__ == '__main__':
    create_document()
